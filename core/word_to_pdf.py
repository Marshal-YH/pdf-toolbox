import docx
import os
import base64
import tempfile, shutil
from core.html_renderer import render_html_to_pdf

def word_to_pdf(input_path: str, output_path: str,
                quality: str = "标准", watermark: str = "") -> int:
    """
    Word 转 PDF（通过 HTML 中间格式）

    Args:
        input_path: .docx 文件路径
        output_path: 输出 PDF 路径
        quality: 质量选项（标准/高清/最小体积）
        watermark: 可选水印文字

    Returns:
        总页数
    """
    # 标准化质量选项
    quality_map = {"标准（推荐）": "标准", "高清": "高清", "最小体积": "最小体积"}
    quality = quality_map.get(quality, "标准")

    doc = docx.Document(input_path)
    html_parts = []

    html_parts.append("""<!DOCTYPE html>
<html><head><meta charset="utf-8">
<style>
    @page { size: A4; margin: 20mm; }
    body { font-family: "Helvetica Neue", "PingFang SC", "Microsoft YaHei", serif;
           font-size: 11pt; line-height: 1.6; color: #1F2937; }
    h1 { font-size: 18pt; font-weight: 700; color: #111827; margin-top: 18pt; margin-bottom: 8pt; }
    h2 { font-size: 14pt; font-weight: 600; color: #1F2937; margin-top: 14pt; margin-bottom: 6pt; }
    h3 { font-size: 12pt; font-weight: 600; color: #374151; margin-top: 10pt; margin-bottom: 4pt; }
    p { margin: 3pt 0; text-indent: 0; }
    table { border-collapse: collapse; width: 100%; margin: 6pt 0; page-break-inside: auto; }
    thead { display: table-header-group; }
    tr { page-break-inside: avoid; }
    th, td { border: 1px solid #D1D5DB; padding: 5pt 7pt; font-size: 9pt; }
    th { background: #F3F4F6; font-weight: 600; text-align: left; }
    img { max-width: 100%; height: auto; margin: 6pt 0; }
    ul, ol { margin: 3pt 0; padding-left: 22pt; }
    li { margin: 2pt 0; }
</style></head><body>""")

    # 处理段落
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            html_parts.append('<p>&nbsp;</p>')
            continue

        style_name = para.style.name if para.style else "Normal"

        # 构建行内样式（加粗/斜体）
        inline_html = ""
        for run in para.runs:
            t = _escape_html(run.text)
            if run.bold and run.italic:
                inline_html += f'<b><i>{t}</i></b>'
            elif run.bold:
                inline_html += f'<b>{t}</b>'
            elif run.italic:
                inline_html += f'<i>{t}</i>'
            else:
                inline_html += t
        if not inline_html:
            inline_html = _escape_html(text)

        if style_name.startswith("Heading 1"):
            html_parts.append(f'<h1>{inline_html}</h1>')
        elif style_name.startswith("Heading 2"):
            html_parts.append(f'<h2>{inline_html}</h2>')
        elif style_name.startswith("Heading 3"):
            html_parts.append(f'<h3>{inline_html}</h3>')
        elif style_name.startswith("List"):
            html_parts.append(f'<li>{inline_html}</li>')
        else:
            html_parts.append(f'<p>{inline_html}</p>')

    # 处理表格
    for table in doc.tables:
        html_parts.append('<table>')
        for i, row in enumerate(table.rows):
            tag = "th" if i == 0 else "td"
            html_parts.append(f'<tr>')
            for cell in row.cells:
                html_parts.append(f'<{tag}>{_escape_html(cell.text)}</{tag}>')
            html_parts.append('</tr>')
        html_parts.append('</table>')

    # 处理图片
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            try:
                image = rel.target_part
                img_data = base64.b64encode(image.blob).decode()
                ext = image.content_type.split("/")[-1]
                if ext in ("jpeg", "png", "gif"):
                    html_parts.append(
                        f'<p><img src="data:image/{ext};base64,{img_data}"></p>'
                    )
            except:
                pass

    html_parts.append('</body></html>')

    # 渲染 PDF — 传递 quality 参数
    full_html = "\n".join(html_parts)
    page_count = render_html_to_pdf(full_html, output_path, quality=quality)

    # 后处理：水印（使用支持中文的字体）
    if watermark:
        _add_watermark(output_path, watermark)

    return page_count


def _add_watermark(path: str, text: str):
    """使用内置 CJK 字体添加水印"""
    import fitz
    doc = fitz.open(path)

    for page in doc:
        w, h = page.rect.width, page.rect.height
        # 使用 PyMuPDF 内置中文字体 china-ss（黑体），不嵌入额外字体文件
        page.insert_text(
            fitz.Point(w / 3, h / 2),
            text,
            fontsize=48,
            fontname="china-ss",
            color=(0.75, 0.75, 0.75),
            overlay=True
        )

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
    tmp_path = tmp.name
    tmp.close()
    doc.save(tmp_path)
    doc.close()
    shutil.move(tmp_path, path)


def _escape_html(text: str) -> str:
    if text is None:
        return ""
    text = str(text)
    text = text.replace("&", "&amp;")
    text = text.replace("<", "&lt;")
    text = text.replace(">", "&gt;")
    text = text.replace('"', "&quot;")
    return text
